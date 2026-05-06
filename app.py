from flask import Flask, request, jsonify, render_template_string, session, send_file
import anthropic
import os
import io
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "bos-deliberations-secret-2026")


def clean_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\*([^\*\n]+)\*', r'\1', text)
    text = re.sub(r'(?m)^\*\s+', '', text)
    text = re.sub(r'(?m)^#+\s+', '', text)
    return text

FREE_LIMIT = 5
usage_by_ip = defaultdict(int)
VALID_CODES = set(
    c.strip().strip('"').strip("'").upper()
    for c in os.environ.get("ACCESS_CODES", "").split(",")
    if c.strip().strip('"').strip("'")
)

SYSTEM_PROMPT = """Tu es un juriste expert en droit public des collectivités territoriales françaises, spécialisé dans la rédaction d'actes administratifs.

Tu rédiges des délibérations de conseil municipal ou communautaire conformes aux exigences légales françaises et à la pratique des collectivités territoriales.

Structure obligatoire d'une délibération :

1. EN-TÊTE
   - Nom de la collectivité en majuscules
   - "CONSEIL MUNICIPAL" ou "CONSEIL COMMUNAUTAIRE" selon le cas
   - "Séance du [date]"
   - "DÉLIBÉRATION N° ____"

2. PRÉSENCE ET QUORUM
   Commencer par la liste des présents sous cette forme exacte :

   "Étaient présents : M./Mme [NOM Prénom], M./Mme [NOM Prénom], M./Mme [NOM Prénom], (...)"
   (Laisser une ligne avec des tirets pour que la commune puisse remplir les noms : "Étaient présents : ____________________________________________")

   Puis les pouvoirs :
   "Avaient donné pouvoir : M./Mme [NOM] avait donné pouvoir à M./Mme [NOM]"
   (Laisser une ligne : "Avaient donné pouvoir : ___________________________________________")

   Puis les absents non représentés :
   "Étaient absents et non représentés : M./Mme [NOM Prénom], (...)"
   (Laisser une ligne : "Étaient absents et non représentés : _____________________________")

   Puis le récapitulatif chiffré :
   "Nombre de membres en exercice : ___"
   "Nombre de membres présents : ___"
   "Nombre de membres ayant donné pouvoir : ___"
   "Quorum atteint : OUI"

3. VISAS (commencer chaque ligne par "VU")
   IMPORTANT sur les visas : ne jamais inventer des numéros d'articles précis. Citer uniquement :
   - Les codes et lois de manière générale ("Vu le Code général des collectivités territoriales")
   - Les grandes divisions connues avec certitude ("notamment le Titre II du Livre II")
   - Les documents locaux concrets fournis dans le contexte (délibérations antérieures, rapports, avis)
   - Toujours terminer les visas par : "Vu les autres textes législatifs et réglementaires applicables en la matière ;"
   - Ajouter en fin de section visas : "[Note : les références législatives précises sont à vérifier sur Légifrance avant présentation en conseil]"

4. EXPOSÉ DES MOTIFS
   "Monsieur/Madame le Maire expose à l'assemblée que..."
   Explication claire du contexte, des enjeux et de la nécessité de la délibération.

5. CONSIDÉRANTS
   "CONSIDÉRANT que [motif 1] ;"
   "CONSIDÉRANT que [motif 2] ;"

6. DISPOSITIF (après la formule "Après en avoir délibéré")
   Décision formulée clairement en utilisant les termes appropriés :
   DECIDE / APPROUVE / AUTORISE / ACCEPTE / DIT / PRÉCISE
   Numéroter les articles si plusieurs décisions.

7. MENTIONS OBLIGATOIRES
   Voies et délais de recours (Tribunal Administratif, 2 mois).
   "Pour extrait conforme, Le Maire, [Signature]"

Rédige une délibération complète, formelle, juridiquement irréprochable. Cite les articles de loi exacts. Utilise le style administratif rigoureux des collectivités françaises.
"""

with open(os.path.join(os.path.dirname(__file__), "index.html"), encoding="utf-8") as f:
    HTML_TEMPLATE = f.read()


def get_client_ip():
    if request.headers.get("X-Forwarded-For"):
        return request.headers.get("X-Forwarded-For").split(",")[0].strip()
    return request.remote_addr


@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/actia-logo.png")
def logo():
    return send_file(os.path.join(os.path.dirname(__file__), "actia-logo.png"), mimetype="image/png")


@app.route("/usage", methods=["GET"])
def get_usage():
    ip = get_client_ip()
    used = usage_by_ip[ip]
    server_key_available = bool(os.environ.get("ANTHROPIC_API_KEY"))
    remaining = max(0, FREE_LIMIT - used) if server_key_available else 0
    return jsonify({
        "remaining": remaining,
        "server_key_available": server_key_available
    })


CHECKLIST_SYSTEM = """Tu es un expert en droit public des collectivités territoriales françaises.
À partir d'une délibération fournie, tu génères une checklist de vérification pratique pour le secrétaire de mairie ou l'agent en charge.

Format de réponse STRICT — deux sections séparées par "---" :

SECTION 1 : Points généraux (toujours présents)
Liste de 4 à 6 points à vérifier pour toute délibération : convocation, quorum, dates, visas, signatures, transmission au contrôle de légalité, etc.

SECTION 2 : Points spécifiques
Liste de 4 à 8 points propres au TYPE de délibération fournie. Sois précis et concret : délais réglementaires, obligations de publicité, pièces justificatives, validations préalables, etc.

Chaque point doit être une phrase courte et actionnable, commençant par un verbe à l'infinitif.
Ne mets pas d'astérisques, pas de markdown. Juste le texte brut avec les listes."""


@app.route("/checklist", methods=["POST"])
def generate_checklist():
    data = request.json
    user_api_key = data.get("api_key", "").strip()
    access_code = data.get("access_code", "").strip()
    deliberation_text = data.get("deliberation_text", "")
    objet = data.get("objet", "")
    type_collectivite = data.get("type_collectivite", "Commune")
    server_key = os.environ.get("ANTHROPIC_API_KEY", "")
    has_valid_code = bool(VALID_CODES) and access_code.upper() in VALID_CODES

    if user_api_key:
        api_key = user_api_key
    elif has_valid_code or server_key:
        api_key = server_key
        if not api_key:
            return jsonify({"error": "Clé API non configurée."}), 500
    else:
        return jsonify({"error": "Clé API manquante."}), 400

    if not deliberation_text:
        return jsonify({"error": "Aucune délibération fournie."}), 400

    prompt = f"""Voici une délibération de type "{objet}" pour une {type_collectivite}.

{deliberation_text[:3000]}

Génère la checklist de vérification selon le format demandé."""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=800,
            system=CHECKLIST_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        parts = raw.split("---")
        generique = parts[0].strip() if len(parts) > 0 else ""
        specifique = parts[1].strip() if len(parts) > 1 else ""

        def parse_items(text):
            lines = [l.strip().lstrip("-•·") .strip() for l in text.split("\n") if l.strip()]
            return [l for l in lines if len(l) > 10 and not l.isupper()]

        return jsonify({
            "generique": parse_items(generique),
            "specifique": parse_items(specifique)
        })
    except Exception as e:
        return jsonify({"error": f"Erreur checklist : {str(e)}"}), 500


@app.route("/validate-code", methods=["POST"])
def validate_code():
    code = request.json.get("code", "").strip().upper()
    valid = bool(VALID_CODES) and code in VALID_CODES
    return jsonify({"valid": valid, "codes_loaded": len(VALID_CODES)})


@app.route("/generate", methods=["POST"])
def generate():
    data = request.json
    user_api_key = data.get("api_key", "").strip()
    access_code = data.get("access_code", "").strip()
    ip = get_client_ip()
    server_key = os.environ.get("ANTHROPIC_API_KEY", "")
    has_valid_code = bool(VALID_CODES) and access_code.upper() in VALID_CODES

    if user_api_key:
        api_key = user_api_key
    elif has_valid_code:
        if not server_key:
            return jsonify({"error": "Clé API serveur non configurée. Contactez le support."}), 500
        api_key = server_key
    elif server_key:
        if usage_by_ip[ip] >= FREE_LIMIT:
            return jsonify({
                "error": f"Vous avez utilisé vos {FREE_LIMIT} générations gratuites. Entrez votre code d'accès abonné ou votre propre clé API Anthropic pour continuer."
            }), 429
        api_key = server_key
        usage_by_ip[ip] += 1
    else:
        return jsonify({"error": "Clé API manquante. Entrez votre clé Anthropic."}), 400

    commune = data.get("commune", "")
    type_collectivite = data.get("type_collectivite", "Commune")
    date_seance = data.get("date_seance", "")
    objet = data.get("objet", "")
    contexte = data.get("contexte", "")
    montant = data.get("montant", "")

    if not objet or not commune:
        return jsonify({"error": "Veuillez renseigner au minimum le nom de la collectivité et l'objet."}), 400

    user_prompt = f"""Rédige une délibération complète avec les informations suivantes :

Collectivité : {commune}
Type de collectivité : {type_collectivite}
Date de séance : {date_seance if date_seance else "À compléter"}
Objet de la délibération : {objet}
Contexte et détails : {contexte if contexte else "Non précisé"}
{f"Montant / budget concerné : {montant}" if montant else ""}

Produis une délibération complète, formelle et juridiquement conforme, prête à être soumise au vote."""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=2500,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        if user_api_key or has_valid_code:
            remaining = None
        else:
            remaining = max(0, FREE_LIMIT - usage_by_ip[ip])
        return jsonify({
            "result": clean_markdown(message.content[0].text),
            "remaining": remaining
        })

    except anthropic.AuthenticationError:
        return jsonify({"error": "Clé API invalide. Vérifiez votre clé sur console.anthropic.com"}), 401
    except anthropic.RateLimitError:
        return jsonify({"error": "Limite de requêtes atteinte. Réessayez dans quelques secondes."}), 429
    except Exception as e:
        return jsonify({"error": f"Erreur inattendue : {str(e)}"}), 500


def build_docx(text, commune, objet, date_seance):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header — commune name
    header = doc.add_paragraph(commune.upper() if commune else "COLLECTIVITÉ")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    # Objet
    if objet:
        sub = doc.add_paragraph(f"Objet : {objet}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(11)
        sub.runs[0].italic = True

    if date_seance:
        d = doc.add_paragraph(f"Séance du {date_seance}")
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        d.runs[0].font.size = Pt(11)

    doc.add_paragraph("")

    # Body — parse lines
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()

        # Detect uppercase header lines (VU, CONSIDÉRANT, DECIDE, APPROUVE...)
        is_header = (
            stripped.isupper() or
            stripped.startswith("VU ") or
            stripped.startswith("CONSIDÉRANT") or
            any(stripped.startswith(w) for w in ["DECIDE", "APPROUVE", "AUTORISE", "ACCEPTE", "DIT ", "PRÉCISE", "Pour extrait"])
        )

        run = p.add_run(stripped)
        run.font.size = Pt(11)
        if is_header:
            run.bold = True

    # Footer disclaimer
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph(
        "⚠️ Ce document est un brouillon généré par ActIA (intelligence artificielle). "
        "Il doit être relu et validé par un professionnel compétent avant toute utilisation officielle. "
        "Les références législatives sont à vérifier sur Légifrance."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].italic = True
    disclaimer.runs[0].font.color.rgb = RGBColor(0x92, 0x40, 0x0e)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/export", methods=["POST"])
def export_docx():
    data = request.json
    text = data.get("text", "")
    commune = data.get("commune", "")
    objet = data.get("objet", "")
    date_seance = data.get("date_seance", "")

    if not text:
        return jsonify({"error": "Aucun contenu à exporter"}), 400

    buf = build_docx(text, commune, objet, date_seance)
    filename = re.sub(r"[^a-z0-9-]", "", objet.lower().replace(" ", "-"))[:40] or "deliberation"
    filename += ".docx"

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") == "development"
    if not debug:
        print("")
        print("  ⚖️  ActIA — Actes administratifs par IA")
        print("  ─────────────────────────────")
        print(f"  Ouvre ton navigateur sur :")
        print(f"  → http://localhost:{port}")
        print("")
    app.run(host="0.0.0.0", port=port, debug=debug)
